from __future__ import annotations

import logging
from pathlib import Path
import re

from docx import Document

from filldoc.core.errors import GenerationError
from filldoc.generator.filename_rules import safe_filename, ensure_unique_path

log = logging.getLogger("filldoc.generator")


def _replace_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    """
    Аккуратная замена по runs, чтобы сохранить форматирование.

    Предполагаем, что:
    - в шаблоне переменные имеют вид {Имя};
    - в mapping ключи — это внутреннее имя без фигурных скобок.
    """
    if not mapping or not paragraph.runs:
        return

    keys = [k for k in mapping.keys() if k]
    if not keys:
        return

    pattern = re.compile(
        "(" + "|".join(r"\{" + re.escape(k) + r"\}" for k in keys) + ")"
    )

    while True:
        runs = list(paragraph.runs)
        full_text = "".join(r.text or "" for r in runs)
        if not full_text:
            return

        m = pattern.search(full_text)
        if not m:
            return

        start, end = m.span()
        placeholder = m.group(0)
        inner = placeholder[1:-1]
        replacement = mapping.get(inner, "")

        positions = []
        pos = 0
        for r in runs:
            t = r.text or ""
            positions.append((r, pos, pos + len(t)))
            pos += len(t)

        overlapping = [
            (r, s, e)
            for (r, s, e) in positions
            if not (e <= start or s >= end)
        ]
        if not overlapping:
            return

        first_run, first_s, first_e = overlapping[0]
        last_run, last_s, last_e = overlapping[-1]

        for r, s, e in overlapping:
            text = r.text or ""
            if r is first_run and r is last_run:
                local_start = max(start - s, 0)
                local_end = min(end - s, len(text))
                before = text[:local_start]
                after = text[local_end:]
                r.text = before + replacement + after
            elif r is first_run:
                local_start = max(start - s, 0)
                before = text[:local_start]
                r.text = before + replacement
            elif r is last_run:
                local_end = min(end - s, len(text))
                after = text[local_end:]
                r.text = after
            else:
                r.text = ""


def _replace_in_table(table, mapping: dict[str, str]) -> None:
    """Обходит таблицу рекурсивно (поддержка вложенных таблиц)."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)
            # Рекурсивно обходим вложенные таблицы
            for nested in cell.tables:
                _replace_in_table(nested, mapping)


def _replace_in_header_footer(hf_part, mapping: dict[str, str]) -> None:
    """Заменяет переменные в колонтитуле (header или footer)."""
    if hf_part.is_linked_to_previous:
        return
    for p in hf_part.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in hf_part.tables:
        _replace_in_table(table, mapping)


def generate_docx_from_template(
    template_path: str,
    output_dir: str,
    output_name: str,
    mapping: dict[str, str],
) -> str:
    try:
        doc = Document(template_path)
    except Exception as e:  # noqa: BLE001
        raise GenerationError(f"Не удалось открыть шаблон для генерации: {e}") from e

    # Основное тело документа
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)

    for table in doc.tables:
        _replace_in_table(table, mapping)

    # Колонтитулы всех секций
    for section in doc.sections:
        _replace_in_header_footer(section.header, mapping)
        _replace_in_header_footer(section.footer, mapping)
        _replace_in_header_footer(section.even_page_header, mapping)
        _replace_in_header_footer(section.even_page_footer, mapping)
        _replace_in_header_footer(section.first_page_header, mapping)
        _replace_in_header_footer(section.first_page_footer, mapping)

    out_dir = Path(output_dir)
    if not out_dir.exists():
        raise GenerationError("Папка выгрузки недоступна или не существует.")

    filename = safe_filename(output_name) + ".docx"
    out_path = ensure_unique_path(out_dir / filename)

    try:
        doc.save(str(out_path))
        log.info("Generated: %s", out_path)
        return str(out_path)
    except Exception as e:  # noqa: BLE001
        raise GenerationError(f"Не удалось сохранить готовый документ: {e}") from e

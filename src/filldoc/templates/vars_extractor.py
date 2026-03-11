from __future__ import annotations

import re
from collections import OrderedDict

from docx import Document

from filldoc.core.errors import TemplateError


_var_re = re.compile(r"\{([^{}]+)\}")


def _extract_from_text(text: str) -> list[str]:
    return [m.group(1).strip() for m in _var_re.finditer(text or "")]


def extract_docx_variables(docx_path: str) -> tuple[list[str], list[str]]:
    """
    MVP: переменная = любой фрагмент в фигурных скобках: `{Переменная}`.
    Собираем по порядку первого появления и без повторов.

    Ограничение MVP: если Word “разбил” `{Переменная}` на разные runs/ячейки,
    мы все равно извлечем ее, т.к. работаем с объединенным текстом параграфа/ячейки.
    """
    try:
        doc = Document(docx_path)
    except Exception as e:  # noqa: BLE001
        raise TemplateError(f"Шаблон поврежден или не удалось прочитать: {e}") from e

    ordered: list[str] = []

    def add_many(items: list[str]) -> None:
        ordered.extend(items)

    for p in doc.paragraphs:
        add_many(_extract_from_text(p.text))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                add_many(_extract_from_text(cell.text))

    uniq = list(OrderedDict((v, True) for v in ordered).keys())
    return ordered, uniq


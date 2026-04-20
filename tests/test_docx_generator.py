from __future__ import annotations

import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from filldoc.generator.docx_generator import generate_docx_from_template


def _save_docx(
    path: Path,
    *,
    paragraph: str | None = None,
    table: str | None = None,
    header: str | None = None,
) -> Path:
    doc = Document()
    if paragraph is not None:
        doc.add_paragraph(paragraph)
    if table is not None:
        t = doc.add_table(rows=1, cols=1)
        t.cell(0, 0).text = table
    if header is not None:
        section = doc.sections[0]
        section.header.is_linked_to_previous = False
        section.header.paragraphs[0].text = header
    doc.save(str(path))
    return path


def test_replaces_variable_in_regular_paragraph(tmp_path: Path) -> None:
    template = _save_docx(tmp_path / "template.docx", paragraph="Client: {client}")

    out_path = generate_docx_from_template(
        str(template),
        str(tmp_path),
        "result",
        {"client": "Acme LLC"},
    )

    doc = Document(out_path)
    assert doc.paragraphs[0].text == "Client: Acme LLC"


def test_replaces_variable_in_table(tmp_path: Path) -> None:
    template = _save_docx(tmp_path / "template.docx", table="Amount: {amount}")

    out_path = generate_docx_from_template(
        str(template),
        str(tmp_path),
        "table-result",
        {"amount": "1000"},
    )

    doc = Document(out_path)
    assert doc.tables[0].cell(0, 0).text == "Amount: 1000"


def test_replaces_variable_in_header(tmp_path: Path) -> None:
    template = _save_docx(tmp_path / "template.docx", header="Case {case_id}")

    out_path = generate_docx_from_template(
        str(template),
        str(tmp_path),
        "header-result",
        {"case_id": "A-42"},
    )

    doc = Document(out_path)
    assert doc.sections[0].header.paragraphs[0].text == "Case A-42"


def test_existing_output_name_gets_numbered_copy(tmp_path: Path) -> None:
    template = _save_docx(tmp_path / "template.docx", paragraph="{name}")
    (tmp_path / "report.docx").write_bytes(b"already exists")

    out_path = generate_docx_from_template(
        str(template),
        str(tmp_path),
        "report",
        {"name": "Generated"},
    )

    assert Path(out_path).name == "report (2).docx"
    assert (tmp_path / "report.docx").read_bytes() == b"already exists"
    assert Document(out_path).paragraphs[0].text == "Generated"
